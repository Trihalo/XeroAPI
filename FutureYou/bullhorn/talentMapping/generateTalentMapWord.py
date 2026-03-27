#!/usr/bin/env python3
"""
generateTalentMapWord.py

Reads an existing Excel talent map (.xlsx) and produces a formatted Word
document (.docx) ready for PDF export.

Layout (every page)
───────────────────
  Page header:  «Client  |  Job Title – Talent Mapping»        [FutureYou logo]
  Body:         Single candidate table; header row repeats on every page.

Usage:
    python3 generateTalentMapWord.py --excel output/FYTalentMap_UnderTheHammer_Mar26.xlsx
    python3 generateTalentMapWord.py --excel output/FYTalentMap_UnderTheHammer_Mar26.xlsx \\
            --client "Under The Hammer" --job-title "Customer Service Manager"
"""

import os
import sys
import argparse
import glob
from datetime import datetime

import openpyxl
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── Brand colours & fonts ────────────────────────────────────────────────────────
NAVY       = RGBColor(0x00, 0x34, 0x64)  # #003464
CORAL      = RGBColor(0xF2, 0x5A, 0x57)  # #F25A57 Salmon
HEADER_HEX = "EEEEEE"                    # Light Grey
WHITE_HEX  = "FFFFFF"
GREY_HEX   = "F7F7F7"                    # subtle alternating row tint

FONT_BODY   = "Trebuchet MS"
FONT_HEADER = "Raleway"

TABLE_HEADERS = ["Name", "Current Company", "Role Title", "Location", "Salary", "LI Profile", "Notes"]
_COL_RATIOS   = [3.5, 3.8, 5.5, 2.8, 2.5, 3.0, 5.6]   # cm; LinkedIn narrowed, Notes wider; sums to 26.7 cm

_LOGO_DEFAULT = os.path.normpath(
    os.path.join(os.path.dirname(__file__),
                 "../../forecastingWebsite/frontend/public/fy.png")
)
_YOU_LOGO_DEFAULT = os.path.normpath(
    os.path.join(os.path.dirname(__file__),
                 "../../forecastingWebsite/frontend/public/fy_you.png")
)


# ── Excel reader ────────────────────────────────────────────────────────────────
def _get_excel_highlight(cell):
    """Return 6-char RRGGBB hex if the cell has a solid RGB fill, else None."""
    try:
        fill = cell.fill
        if fill.fill_type != "solid":
            return None
        fg = fill.fgColor
        if fg.type != "rgb":
            return None
        rgba = fg.rgb  # AARRGGBB, e.g. "FF92D050" (green)
        if not rgba or rgba in ("00000000", "FF000000"):
            return None
        return rgba[2:]  # strip alpha → RRGGBB
    except Exception:
        return None


def read_excel(path):
    """Return list of row dicts from the talent map Excel, skipping the header row."""
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    candidates = []
    for row in ws.iter_rows(min_row=2, values_only=False):
        # Skip completely empty rows
        values = [cell.value for cell in row]
        if not any(values):
            continue
        # LinkedIn may be stored as a hyperlink rather than a plain value
        li_cell  = row[5] if len(row) > 5 else None
        li_value = ""
        if li_cell:
            if li_cell.hyperlink:
                li_value = li_cell.hyperlink.target or li_cell.value or ""
            else:
                li_value = li_cell.value or ""

        candidates.append({
            "name":      str(values[0] or "").strip(),
            "company":   str(values[1] or "").strip(),
            "role":      str(values[2] or "").strip(),
            "location":  str(values[3] or "").strip(),
            "salary":    str(values[4] or "").strip(),
            "linkedin":  str(li_value).strip(),
            "notes":     str(values[6] or "").strip() if len(values) > 6 else "",
            "highlight": _get_excel_highlight(row[0]),
        })
    return candidates


# ── Word XML helpers ─────────────────────────────────────────────────────────────
def _set_cell_fill(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_table_borders(tbl_el, val="none"):
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), val)
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def _clear_outer_table_borders(tbl_el):
    """Remove only the outer perimeter borders, keeping inner gridlines intact."""
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def _mark_row_as_header(row):
    """Causes this table row to repeat at the top of each new page."""
    trPr      = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _add_row_bottom_border(row, color="D0D5DD", sz="4"):
    for cell in row.cells:
        tcPr      = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        bottom    = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    sz)
        bottom.set(qn("w:color"), color)
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


# ── Hyperlink helper ─────────────────────────────────────────────────────────────
def _add_hyperlink(paragraph, url, text):
    """Insert a clickable hyperlink run into a paragraph."""
    r_id      = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr    = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_BODY)
    rFonts.set(qn("w:hAnsi"), FONT_BODY)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")   # 9pt × 2
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")   # standard hyperlink blue
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run_el.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run_el.append(t)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


# ── Page header ──────────────────────────────────────────────────────────────────
def _build_page_header(doc, title_text, logo_path, usable_width_inches):
    section = doc.sections[0]
    header  = section.header

    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header.add_table(rows=1, cols=2, width=Inches(usable_width_inches))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(tbl._tbl, val="none")

    # Left: title
    left = tbl.cell(0, 0)
    left.width = Inches(usable_width_inches * 0.68)
    lp  = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = lp.add_run(title_text)
    run.font.name      = FONT_HEADER
    run.font.size      = Pt(13)
    run.font.color.rgb = CORAL

    # Right: logo
    right = tbl.cell(0, 1)
    right.width = Inches(usable_width_inches * 0.32)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if logo_path and os.path.exists(logo_path):
        rp.add_run().add_picture(logo_path, height=Inches(0.42))
    else:
        r2 = rp.add_run("FutureYou")
        r2.font.color.rgb = CORAL
        r2.font.size      = Pt(14)


# ── Page footer ──────────────────────────────────────────────────────────────────
def _build_page_footer(doc, you_logo_path):
    """Place the 'You' wordmark as a floating anchor pinned to bottom-left of every page."""
    if not you_logo_path or not os.path.exists(you_logo_path):
        return

    section = doc.sections[0]
    footer  = section.footer

    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    p = footer.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

    # Embed via add_picture so python-docx wires the image relationship
    run = p.add_run()
    run.add_picture(you_logo_path, height=Cm(3.5))

    # Locate the <wp:inline> element that add_picture just created
    drawing_el = run._r.find(qn("w:drawing"))
    inline_el  = drawing_el.find(qn("wp:inline"))

    # Read actual image extents (EMU) from the inline element
    extent = inline_el.find(qn("wp:extent"))
    cx     = int(extent.get("cx"))
    cy     = int(extent.get("cy"))

    # Absolute position on the page (all values in EMU; 1 cm = 360 000 EMU)
    page_h_emu = int(section.page_height)          # e.g. 7 560 000 for A4 landscape
    x_emu      = 0                                  # hard against the left page edge
    y_emu      = page_h_emu - cy - int(Cm(0.15))   # ~0.15 cm above the very bottom

    # ── Build <wp:anchor> ────────────────────────────────────────────────────────
    anchor = OxmlElement("wp:anchor")
    anchor.set("distT",          "0")
    anchor.set("distB",          "0")
    anchor.set("distL",          "0")
    anchor.set("distR",          "0")
    anchor.set("simplePos",      "0")
    anchor.set("relativeHeight", "251658240")
    anchor.set("behindDoc",      "0")
    anchor.set("locked",         "0")
    anchor.set("layoutInCell",   "1")
    anchor.set("allowOverlap",   "1")

    sp = OxmlElement("wp:simplePos")
    sp.set("x", "0"); sp.set("y", "0")
    anchor.append(sp)

    posH = OxmlElement("wp:positionH")
    posH.set("relativeFrom", "page")
    ph_off = OxmlElement("wp:posOffset")
    ph_off.text = str(x_emu)
    posH.append(ph_off)
    anchor.append(posH)

    posV = OxmlElement("wp:positionV")
    posV.set("relativeFrom", "page")
    pv_off = OxmlElement("wp:posOffset")
    pv_off.text = str(y_emu)
    posV.append(pv_off)
    anchor.append(posV)

    ext2 = OxmlElement("wp:extent")
    ext2.set("cx", str(cx)); ext2.set("cy", str(cy))
    anchor.append(ext2)

    ee = OxmlElement("wp:effectExtent")
    ee.set("l", "0"); ee.set("t", "0"); ee.set("r", "0"); ee.set("b", "0")
    anchor.append(ee)

    anchor.append(OxmlElement("wp:wrapNone"))

    # Copy <wp:docPr> and <wp:cNvGraphicFramePr> from the original inline
    for tag in (qn("wp:docPr"), qn("wp:cNvGraphicFramePr")):
        el = inline_el.find(tag)
        if el is not None:
            anchor.append(el)

    # Copy <a:graphic> (the actual image data) — match by tag suffix to be safe
    for child in inline_el:
        if child.tag.endswith("}graphic"):
            anchor.append(child)
            break

    # Swap inline → anchor
    drawing_el.remove(inline_el)
    drawing_el.append(anchor)


# ── Document builder ─────────────────────────────────────────────────────────────
def build_word(candidates, client_corp, job_title, logo_path, you_logo_path=None):
    doc = Document()

    # Landscape A4
    section               = doc.sections[0]
    section.page_width    = Cm(29.7)
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.8)
    section.bottom_margin    = Cm(4.5)   # reserves space so table never overlaps "You"
    section.header_distance  = Cm(1.0)
    section.footer_distance  = Cm(0.2)  # "You" sits tight in the bottom-left corner

    usable_w_cm     = 29.7 - 1.5 - 1.5
    usable_w_inches = usable_w_cm / 2.54

    _build_page_header(doc, f"{client_corp}  |  {job_title} – Talent Mapping",
                       logo_path, usable_w_inches)
    _build_page_footer(doc, you_logo_path)

    total_ratio = sum(_COL_RATIOS)
    col_widths  = [Cm(r * usable_w_cm / total_ratio) for r in _COL_RATIOS]

    table           = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = "Table Grid"
    _clear_outer_table_borders(table._tbl)   # keep inner grid, remove perimeter

    # Header row
    hdr_row        = table.rows[0]
    hdr_row.height = Cm(0.85)
    for col_idx, (hdr_text, width) in enumerate(zip(TABLE_HEADERS, col_widths)):
        cell = hdr_row.cells[col_idx]
        cell.width              = width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(hdr_text)
        run.font.name      = FONT_HEADER
        run.font.size      = Pt(9)
        run.font.color.rgb = NAVY
    _add_row_bottom_border(hdr_row, color="000000", sz="8")
    _mark_row_as_header(hdr_row)

    # Data rows — bottom border on every row except the last
    candidates = list(candidates)
    for idx, c in enumerate(candidates):
        row        = table.add_row()
        row.height = Cm(0.75)
        if idx < len(candidates) - 1:
            _add_row_bottom_border(row)

        values = [c["name"], c["company"], c["role"], c["location"],
                  c["salary"], c["linkedin"], c["notes"]]
        for col_idx, (val, width) in enumerate(zip(values, col_widths)):
            cell = row.cells[col_idx]
            cell.width              = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Carry over Excel highlight on the name cell
            if col_idx == 0 and c.get("highlight"):
                _set_cell_fill(cell, c["highlight"])
            p = cell.paragraphs[0]
            # LinkedIn column — clickable hyperlink
            if col_idx == 5 and val and val.startswith("http"):
                _add_hyperlink(p, val, val)
            else:
                run = p.add_run(val or "")
                run.font.name = FONT_BODY
                run.font.size = Pt(9)

    return doc


# ── Main ─────────────────────────────────────────────────────────────────────────
def main():
    # Find the most recent xlsx in the output folder as a sensible default
    output_dir    = os.path.join(os.path.dirname(__file__), "output")
    default_excel = ""
    xlsx_files = sorted(glob.glob(os.path.join(output_dir, "*.xlsx")), key=os.path.getmtime, reverse=True)
    if xlsx_files:
        default_excel = xlsx_files[0]

    parser = argparse.ArgumentParser(
        description="Generate a formatted Word talent map from an existing Excel file."
    )
    parser.add_argument("--excel",     default=default_excel,
                        help=f"Path to the Excel talent map (default: most recent in output/)")
    parser.add_argument("--client",    default=None,
                        help="Client/company name for the document header")
    parser.add_argument("--job-title", default=None,
                        help="Job title for the document header")
    parser.add_argument("--logo",      default=_LOGO_DEFAULT,
                        help="Path to FutureYou logo PNG")
    parser.add_argument("--you-logo",  default=_YOU_LOGO_DEFAULT,
                        help="Path to the 'You' wordmark PNG for the page footer")
    parser.add_argument("--output",    default=output_dir,
                        help="Output directory for the .docx file")
    args = parser.parse_args()

    if not args.excel or not os.path.exists(args.excel):
        sys.exit(f"Excel file not found: {args.excel!r}\nPass --excel path/to/file.xlsx")

    logo_path = args.logo if os.path.exists(args.logo) else None
    if not logo_path:
        print(f"  [Logo not found at {args.logo} — using text fallback]")

    you_logo_path = args.you_logo if os.path.exists(args.you_logo) else None
    if not you_logo_path:
        print(f"  [You logo not found at {args.you_logo} — footer watermark skipped]")

    print(f"Reading Excel: {args.excel}")
    candidates = read_excel(args.excel)
    print(f"  {len(candidates)} candidates loaded")

    client_corp = args.client    or input("Client company name for header: ").strip()
    job_title   = args.job_title or input("Job title for header: ").strip()

    os.makedirs(args.output, exist_ok=True)
    date_str   = datetime.today().strftime("%b%y")
    safe_corp  = client_corp.replace("/", "-").replace(" ", "")
    safe_title = "".join(c for c in job_title if c.isalnum() or c in " -")[:30].strip().replace(" ", "")
    out_path   = os.path.join(args.output, f"FYTalentMap_{safe_corp}_{safe_title}_{date_str}.docx")

    doc = build_word(candidates, client_corp, job_title, logo_path, you_logo_path)
    doc.save(out_path)
    print(f"  → Saved: {out_path}")


if __name__ == "__main__":
    main()
